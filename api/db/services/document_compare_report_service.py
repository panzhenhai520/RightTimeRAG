#
#  Copyright 2026 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

from __future__ import annotations

import json
import os
import re
from copy import deepcopy
from datetime import datetime, timezone
from io import BytesIO
from typing import Any


REPORT_MIME_TYPES = {
    "markdown": "text/markdown",
    "md": "text/markdown",
    "json": "application/json",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


class DocumentCompareReportService:
    """Build auditable reports from document compare results."""

    @classmethod
    def build_report(
        cls,
        *,
        title: str = "文档比对报告",
        files: Any = None,
        documents: Any = None,
        diff: Any = None,
        table_diff: Any = None,
        matches: Any = None,
        conflicts: Any = None,
        missing_requirements: Any = None,
        risk_points: Any = None,
        audit: Any = None,
        run_id: str = "",
        agent_id: str = "",
        generated_by: str = "DocumentCompareReportComposer",
    ) -> dict[str, Any]:
        diff_data = cls._as_dict(diff)
        table_diff_data = cls._as_dict(table_diff)
        match_data = cls._as_dict(matches)
        conflict_data = cls._as_dict(conflicts)
        match_items = cls._items_from(match_data or matches, "matches")
        conflict_items = cls._items_from(conflict_data or conflicts, "conflicts")
        missing_items = cls._items_from(missing_requirements, "missing_requirements")
        if not missing_items:
            missing_items = cls._items_from(conflict_data, "missing_requirements")
        risk_items = cls._items_from(risk_points, "risk_points")
        file_items = cls._collect_files(files, documents)
        audit_payload = cls._build_audit(
            audit=audit,
            documents=documents,
            run_id=run_id,
            agent_id=agent_id,
            generated_by=generated_by,
        )
        report = {
            "schema_version": 1,
            "title": title or "文档比对报告",
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "run_id": run_id,
            "agent_id": agent_id,
            "files": file_items,
            "parse_summary": cls._parse_summary(documents),
            "diff_summary": cls._diff_summary(diff_data),
            "table_diff_summary": cls._diff_summary(table_diff_data),
            "matches": match_items,
            "conflicts": conflict_items,
            "missing_requirements": missing_items,
            "risk_points": risk_items,
            "risk_level": cls._risk_level(conflict_items, missing_items, risk_items),
            "references": cls._collect_references([diff_data, table_diff_data, match_items, conflict_items, missing_items, risk_items]),
            "audit": audit_payload,
        }
        report["summary"] = cls._summary(report)
        return report

    @classmethod
    def render_markdown(cls, report: dict[str, Any]) -> str:
        lines = [
            f"# {report.get('title') or '文档比对报告'}",
            "",
            f"- run_id: {report.get('run_id') or ''}",
            f"- agent_id: {report.get('agent_id') or ''}",
            f"- risk_level: {report.get('risk_level') or 'none'}",
            f"- generated_at: {report.get('generated_at') or ''}",
            "",
            "## 文件清单",
            "",
            "| 文件 | 路径 | 类型 | 审计 |",
            "| --- | --- | --- | --- |",
        ]
        for item in report.get("files") or []:
            lines.append(
                "| "
                + " | ".join(
                    [
                        cls._md_cell(item.get("filename") or item.get("name") or ""),
                        cls._md_cell(item.get("source_path") or item.get("path") or item.get("relative_path") or ""),
                        cls._md_cell(item.get("mime_type") or item.get("type") or ""),
                        cls._md_cell(item.get("audit_id") or ""),
                    ]
                )
                + " |"
            )
        if not report.get("files"):
            lines.append("| 无 |  |  |  |")
        lines.extend(["", "## 解析摘要", ""])
        parse_summary = report.get("parse_summary") or {}
        lines.extend(
            [
                f"- document_count: {parse_summary.get('document_count', 0)}",
                f"- paragraph_count: {parse_summary.get('paragraph_count', 0)}",
                f"- line_count: {parse_summary.get('line_count', 0)}",
                f"- table_count: {parse_summary.get('table_count', 0)}",
                "",
                "## 精确 Diff 摘要",
                "",
            ]
        )
        for key, value in (report.get("diff_summary") or {}).items():
            lines.append(f"- {key}: {value}")
        if report.get("table_diff_summary"):
            lines.extend(["", "## 表格 Diff 摘要", ""])
            for key, value in (report.get("table_diff_summary") or {}).items():
                lines.append(f"- {key}: {value}")
        lines.extend(["", "## 条款匹配矩阵", "", "| A | B | 关系 | 置信度 | 证据 |", "| --- | --- | --- | --- | --- |"])
        for item in (report.get("matches") or [])[:120]:
            left = item.get("left_item") or {}
            right = item.get("right_item") or {}
            evidence = item.get("evidence") or {}
            lines.append(
                "| "
                + " | ".join(
                    [
                        cls._md_cell(cls._short_text(left)),
                        cls._md_cell(cls._short_text(right)),
                        cls._md_cell(item.get("relation") or ""),
                        cls._md_cell(item.get("score") or ""),
                        cls._md_cell(cls._evidence_text(evidence)),
                    ]
                )
                + " |"
            )
        if not report.get("matches"):
            lines.append("| 无 | 无 |  |  |  |")
        lines.extend(["", "## 冲突清单", "", "| 严重级别 | 原因 | 标准条款 | 目标条款 | 证据 |", "| --- | --- | --- | --- | --- |"])
        for item in report.get("conflicts") or []:
            lines.append(
                "| "
                + " | ".join(
                    [
                        cls._md_cell(item.get("severity") or ""),
                        cls._md_cell(item.get("reason") or item.get("reason_code") or ""),
                        cls._md_cell(cls._short_text(item.get("standard_item") or {})),
                        cls._md_cell(cls._short_text(item.get("target_item") or {})),
                        cls._md_cell(cls._evidence_text(item.get("evidence") or {})),
                    ]
                )
                + " |"
            )
        if not report.get("conflicts"):
            lines.append("| 无 |  |  |  |  |")
        lines.extend(["", "## 缺失项", "", "| 条款 | 证据 |", "| --- | --- |"])
        for item in report.get("missing_requirements") or []:
            left = item.get("left_item") or item
            lines.append("| " + " | ".join([cls._md_cell(cls._short_text(left)), cls._md_cell(cls._evidence_text(item.get("evidence") or {}))]) + " |")
        if not report.get("missing_requirements"):
            lines.append("| 无 |  |")
        lines.extend(["", "## 审计", ""])
        audit = report.get("audit") or {}
        lines.append(f"- deterministic_steps: {', '.join(audit.get('deterministic_steps') or [])}")
        lines.append(f"- llm_steps: {', '.join(audit.get('llm_steps') or [])}")
        lines.append(f"- report_artifact_count: {len(audit.get('report_artifacts') or [])}")
        return "\n".join(lines).strip() + "\n"

    @classmethod
    def render_json(cls, report: dict[str, Any]) -> str:
        return json.dumps(report, ensure_ascii=False, indent=2)

    @classmethod
    def render_bytes(cls, report: dict[str, Any], output_format: str) -> tuple[bytes, str]:
        fmt = cls.normalize_format(output_format)
        if fmt == "markdown":
            return cls.render_markdown(report).encode("utf-8"), REPORT_MIME_TYPES[fmt]
        if fmt == "json":
            return cls.render_json(report).encode("utf-8"), REPORT_MIME_TYPES[fmt]
        if fmt == "docx":
            return cls.render_docx(report), REPORT_MIME_TYPES[fmt]
        if fmt == "xlsx":
            return cls.render_xlsx(report), REPORT_MIME_TYPES[fmt]
        raise ValueError(f"Unsupported report format: {output_format}")

    @classmethod
    def render_docx(cls, report: dict[str, Any]) -> bytes:
        from docx import Document

        document = Document()
        document.add_heading(report.get("title") or "文档比对报告", level=1)
        document.add_paragraph(f"run_id: {report.get('run_id') or ''}")
        document.add_paragraph(f"risk_level: {report.get('risk_level') or 'none'}")
        document.add_heading("文件清单", level=2)
        for item in report.get("files") or []:
            document.add_paragraph(f"{item.get('filename') or item.get('name') or ''} - {item.get('source_path') or item.get('path') or ''}")
        document.add_heading("精确 Diff 摘要", level=2)
        for key, value in (report.get("diff_summary") or {}).items():
            document.add_paragraph(f"{key}: {value}")
        document.add_heading("冲突清单", level=2)
        for item in report.get("conflicts") or []:
            document.add_paragraph(f"[{item.get('severity')}] {item.get('reason') or item.get('reason_code')}")
            document.add_paragraph(cls._evidence_text(item.get("evidence") or {}))
        document.add_heading("缺失项", level=2)
        for item in report.get("missing_requirements") or []:
            document.add_paragraph(cls._short_text(item.get("left_item") or item))
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @classmethod
    def render_xlsx(cls, report: dict[str, Any]) -> bytes:
        from openpyxl import Workbook

        wb = Workbook()
        summary = wb.active
        summary.title = "Summary"
        cls._write_rows(
            summary,
            [
                ["title", report.get("title")],
                ["run_id", report.get("run_id")],
                ["agent_id", report.get("agent_id")],
                ["risk_level", report.get("risk_level")],
                ["summary", report.get("summary")],
            ],
        )
        cls._write_rows(wb.create_sheet("Files"), [["filename", "path", "mime_type", "audit_id"]] + [[item.get("filename"), item.get("source_path") or item.get("path"), item.get("mime_type"), item.get("audit_id")] for item in report.get("files") or []])
        cls._write_rows(wb.create_sheet("Diff"), [["key", "value"]] + [[key, value] for key, value in (report.get("diff_summary") or {}).items()])
        cls._write_rows(
            wb.create_sheet("Matches"),
            [["relation", "score", "left", "right", "evidence"]]
            + [[item.get("relation"), item.get("score"), cls._short_text(item.get("left_item") or {}), cls._short_text(item.get("right_item") or {}), cls._evidence_text(item.get("evidence") or {})] for item in report.get("matches") or []],
        )
        cls._write_rows(
            wb.create_sheet("Conflicts"),
            [["severity", "reason_code", "reason", "standard", "target", "evidence"]]
            + [[item.get("severity"), item.get("reason_code"), item.get("reason"), cls._short_text(item.get("standard_item") or {}), cls._short_text(item.get("target_item") or {}), cls._evidence_text(item.get("evidence") or {})] for item in report.get("conflicts") or []],
        )
        cls._write_rows(
            wb.create_sheet("Missing"),
            [["text", "evidence"]]
            + [[cls._short_text(item.get("left_item") or item), cls._evidence_text(item.get("evidence") or {})] for item in report.get("missing_requirements") or []],
        )
        cls._write_rows(
            wb.create_sheet("Audit"),
            [["key", "value"]]
            + [[key, json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else value] for key, value in (report.get("audit") or {}).items()],
        )
        buffer = BytesIO()
        wb.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def normalize_format(value: str) -> str:
        fmt = str(value or "markdown").strip().lower()
        if fmt == "md":
            return "markdown"
        if fmt not in {"markdown", "json", "docx", "xlsx"}:
            raise ValueError(f"Unsupported report format: {value}")
        return fmt

    @staticmethod
    def filename(base: str, output_format: str) -> str:
        fmt = DocumentCompareReportService.normalize_format(output_format)
        ext = "md" if fmt == "markdown" else fmt
        safe = str(base or "document_compare_report").strip().replace("\\", " ").replace("/", " ")
        safe = "".join(ch for ch in safe if ch >= " " and ch not in '<>:"|?*#%')
        safe = safe.strip(" .")[:160] or "document_compare_report"
        return f"{safe}.{ext}"

    @staticmethod
    def _write_rows(sheet, rows: list[list[Any]]) -> None:
        for row in rows:
            sheet.append(row)

    @classmethod
    def _collect_files(cls, files: Any, documents: Any) -> list[dict[str, Any]]:
        result = []
        seen = set()
        for item in cls._as_list(files):
            if isinstance(item, dict):
                result.append(deepcopy(item))
        for document in cls._as_list(documents):
            if not isinstance(document, dict):
                continue
            file_info = document.get("file") if isinstance(document.get("file"), dict) else {}
            audit = document.get("audit") if isinstance(document.get("audit"), dict) else {}
            result.append(
                {
                    "document_id": document.get("document_id"),
                    "filename": document.get("filename") or file_info.get("name") or file_info.get("filename"),
                    "source_path": document.get("source_path") or file_info.get("path") or file_info.get("relative_path"),
                    "mime_type": document.get("mime_type") or file_info.get("mime_type"),
                    "audit_id": audit.get("audit_id"),
                }
            )
        normalized = []
        for item in result:
            key = (item.get("document_id"), item.get("filename"), item.get("source_path") or item.get("path"))
            if key in seen:
                continue
            seen.add(key)
            normalized.append(item)
        return normalized

    @classmethod
    def _parse_summary(cls, documents: Any) -> dict[str, int]:
        docs = [item for item in cls._as_list(documents) if isinstance(item, dict)]
        return {
            "document_count": len(docs),
            "paragraph_count": sum(len(item.get("paragraphs") or []) for item in docs),
            "line_count": sum(len(item.get("lines") or []) for item in docs),
            "table_count": sum(len(item.get("tables") or []) for item in docs),
            "chunk_count": sum(len(item.get("chunks") or []) for item in docs),
        }

    @staticmethod
    def _diff_summary(value: dict[str, Any]) -> dict[str, Any]:
        summary = value.get("summary") if isinstance(value, dict) else {}
        return deepcopy(summary) if isinstance(summary, dict) else {}

    @classmethod
    def _build_audit(cls, *, audit: Any, documents: Any, run_id: str, agent_id: str, generated_by: str) -> dict[str, Any]:
        file_access = []
        for document in cls._as_list(documents):
            if isinstance(document, dict) and isinstance(document.get("audit"), dict):
                file_access.append(deepcopy(document["audit"]))
        extra = cls._as_dict(audit)
        return {
            "schema_version": 1,
            "run_id": run_id,
            "agent_id": agent_id,
            "generated_by": generated_by,
            "file_access": file_access,
            "node_runs": extra.get("node_runs", []),
            "model_judgements": extra.get("model_judgements", []),
            "deterministic_steps": extra.get("deterministic_steps") or ["document_diff", "table_diff", "semantic_compare", "conflict_detection"],
            "llm_steps": extra.get("llm_steps") or [],
            "report_artifacts": extra.get("report_artifacts") or [],
        }

    @staticmethod
    def _risk_level(conflicts: list[dict[str, Any]], missing: list[dict[str, Any]], risk_points: list[dict[str, Any]]) -> str:
        severities = {str(item.get("severity") or "").lower() for item in conflicts + risk_points if isinstance(item, dict)}
        if "high" in severities:
            return "high"
        if conflicts or "medium" in severities:
            return "medium"
        if missing:
            return "low"
        return "none"

    @staticmethod
    def _summary(report: dict[str, Any]) -> str:
        return (
            f"{len(report.get('files') or [])} file(s), "
            f"{len(report.get('matches') or [])} match item(s), "
            f"{len(report.get('conflicts') or [])} conflict(s), "
            f"{len(report.get('missing_requirements') or [])} missing requirement(s)."
        )

    @classmethod
    def _collect_references(cls, values: Any) -> list[dict[str, Any]]:
        references = []
        seen = set()

        def visit(value: Any):
            if isinstance(value, dict):
                if value.get("source_ref"):
                    source_ref = str(value.get("source_ref"))
                    if source_ref not in seen:
                        seen.add(source_ref)
                        references.append({"source_ref": source_ref, "page": value.get("page"), "line_start": value.get("line_start"), "line_end": value.get("line_end")})
                for key in ("evidence", "left", "right", "a_evidence", "b_evidence", "standard_item", "target_item", "left_item", "right_item", "references"):
                    if key in value:
                        visit(value[key])
            elif isinstance(value, list):
                for item in value:
                    visit(item)

        visit(values)
        return references

    @classmethod
    def _items_from(cls, value: Any, key: str) -> list[dict[str, Any]]:
        parsed = cls._parse_json_like(value, value)
        if isinstance(parsed, dict):
            items = parsed.get(key)
            if isinstance(items, list):
                return [deepcopy(item) for item in items if isinstance(item, dict)]
            if key == "matches" and isinstance(parsed.get("items"), list):
                return [deepcopy(item) for item in parsed["items"] if isinstance(item, dict)]
            return []
        if isinstance(parsed, list):
            return [deepcopy(item) for item in parsed if isinstance(item, dict)]
        return []

    @staticmethod
    def _as_dict(value: Any) -> dict[str, Any]:
        value = DocumentCompareReportService._parse_json_like(value, value)
        return deepcopy(value) if isinstance(value, dict) else {}

    @staticmethod
    def _as_list(value: Any) -> list[Any]:
        value = DocumentCompareReportService._parse_json_like(value, value)
        if value is None or value == "":
            return []
        if isinstance(value, list):
            return deepcopy(value)
        if isinstance(value, dict):
            return [deepcopy(value)]
        return []

    @staticmethod
    def _parse_json_like(value: Any, default: Any = None) -> Any:
        if isinstance(value, str):
            text = value.strip()
            if not text:
                return default
            try:
                return json.loads(text)
            except Exception:
                return default
        return value if value is not None else default

    @staticmethod
    def _short_text(value: Any, limit: int = 120) -> str:
        if isinstance(value, dict):
            text = value.get("text") or value.get("content") or value.get("requirement") or value.get("claim") or ""
        else:
            text = str(value or "")
        text = re.sub(r"\s+", " ", str(text)).strip()
        return text[:limit] + ("..." if len(text) > limit else "")

    @staticmethod
    def _evidence_text(value: Any) -> str:
        refs = []
        if isinstance(value, dict):
            for key in ("source_ref", "a_evidence", "b_evidence"):
                if isinstance(value.get(key), str) and value.get(key):
                    refs.append(value[key])
            for side in ("left", "right"):
                side_value = value.get(side)
                if isinstance(side_value, dict) and side_value.get("source_ref"):
                    refs.append(side_value["source_ref"])
        if isinstance(value, list):
            for item in value:
                if isinstance(item, dict) and item.get("source_ref"):
                    refs.append(item["source_ref"])
        return "; ".join(str(item) for item in refs if item)

    @staticmethod
    def _md_cell(value: Any) -> str:
        text = str(value if value is not None else "")
        text = text.replace("\n", " ").replace("|", "\\|")
        return text
